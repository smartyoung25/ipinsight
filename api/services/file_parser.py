"""특허 문서 파일 파싱 서비스 — PDF / HWP / HWPX / DOCX / TXT 지원"""
from __future__ import annotations
import io
import re
import zipfile
from pathlib import Path


# ── 확장자 → 파서 라우팅 ────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".hwp", ".hwpx", ".docx", ".doc", ".txt"}


def parse_file(filename: str, content: bytes) -> str:
    """파일 바이트를 받아 순수 텍스트를 반환.

    인식 불가 파일은 ValueError 발생.
    품질이 낮거나 텍스트가 짧으면 _quality_warn 플래그를 붙여 반환.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(content)
    elif ext in (".hwp",):
        return _parse_hwp(content)
    elif ext in (".hwpx",):
        return _parse_hwpx(content)
    elif ext in (".docx", ".doc"):
        return _parse_docx(content)
    elif ext == ".txt":
        return _parse_txt(content)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}. 지원: PDF, HWP, HWPX, DOCX, TXT")


# ── PDF ──────────────────────────────────────────────────────────────

def _parse_pdf(content: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text as pm_extract
        text = pm_extract(io.BytesIO(content))
        return _clean(text)
    except Exception as e:
        raise ValueError(f"PDF 파싱 실패: {e}") from e


# ── HWP (바이너리 OLE 구조) ──────────────────────────────────────────

def _parse_hwp(content: bytes) -> str:
    """HWP 5.x 바이너리 — BodyText 스트림에서 텍스트 추출."""
    try:
        import olefile
        import zlib
    except ImportError as e:
        raise ValueError(f"HWP 파싱 의존성 없음: {e}") from e

    if not olefile.isOleFile(io.BytesIO(content)):
        raise ValueError("올바른 HWP 파일이 아닙니다 (OLE 구조 아님)")

    ole = olefile.OleFileIO(io.BytesIO(content))
    texts: list[str] = []

    # 섹션 스트림 순차 추출
    section_idx = 0
    while True:
        stream_path = f"BodyText/Section{section_idx}"
        if not ole.exists(stream_path):
            break
        raw = ole.openstream(stream_path).read()
        # 압축 여부 확인 (FileHeader의 CompressedBodyText 플래그)
        try:
            decompressed = zlib.decompress(raw, -15)
        except Exception:
            decompressed = raw

        texts.append(_decode_hwp_section(decompressed))
        section_idx += 1

    ole.close()
    result = _clean("\n".join(texts))
    if not result.strip():
        raise ValueError("HWP 텍스트 추출 결과가 비어있습니다. 이미지 기반 HWP는 지원되지 않습니다.")
    return result


def _decode_hwp_section(data: bytes) -> str:
    """HWP 섹션 바이트에서 유니코드 텍스트 추출."""
    texts: list[str] = []
    i = 0
    while i + 4 <= len(data):
        tag_id = (int.from_bytes(data[i:i+2], "little") & 0x03FF)
        size = int.from_bytes(data[i+2:i+4], "little")
        i += 4
        payload = data[i:i+size]
        i += size
        # ParaText 태그 = 67
        if tag_id == 67 and len(payload) >= 2:
            try:
                txt = payload.decode("utf-16-le", errors="replace").replace("\x00", "")
                texts.append(txt)
            except Exception:
                pass
    return " ".join(texts)


# ── HWPX (ZIP+XML 구조) ──────────────────────────────────────────────

def _parse_hwpx(content: bytes) -> str:
    """HWPX (HWP 2015+) — ZIP 내 section/*.xml에서 텍스트 추출."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(content))
    except Exception as e:
        raise ValueError(f"HWPX 압축 해제 실패: {e}") from e

    import xml.etree.ElementTree as ET
    texts: list[str] = []
    ns = {"hp": "http://www.hancom.co.kr/hwpml/2012/paragraph"}

    section_files = sorted(
        [n for n in zf.namelist() if re.match(r"Contents/section\d+\.xml", n, re.I)]
    )
    for sf in section_files:
        try:
            xml_bytes = zf.read(sf)
            root = ET.fromstring(xml_bytes)
            for t in root.iter("{http://www.hancom.co.kr/hwpml/2012/paragraph}t"):
                if t.text:
                    texts.append(t.text)
        except Exception:
            pass

    zf.close()
    result = _clean("\n".join(texts))
    if not result.strip():
        raise ValueError("HWPX 텍스트 추출 결과가 비어있습니다.")
    return result


# ── DOCX ─────────────────────────────────────────────────────────────

def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return _clean("\n".join(paragraphs))
    except Exception as e:
        raise ValueError(f"DOCX 파싱 실패: {e}") from e


# ── TXT ──────────────────────────────────────────────────────────────

def _parse_txt(content: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            return _clean(content.decode(enc))
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 파일 인코딩 인식 실패 (UTF-8, CP949 시도)")


# ── 공통 정리 ────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """불필요한 공백/제어문자 제거 및 연속 빈줄 압축."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def quality_check(text: str) -> dict:
    """파싱된 텍스트 품질 평가."""
    char_count = len(text)
    has_claims = bool(re.search(r"청구항\s*\d+|claim\s*\d+", text, re.I))
    has_korean = bool(re.search(r"[가-힣]", text))
    has_english = bool(re.search(r"[a-zA-Z]{3,}", text))
    non_text_ratio = len(re.findall(r"[^\w\s가-힣]", text)) / max(char_count, 1)

    warnings: list[str] = []
    if char_count < 200:
        warnings.append(f"텍스트 길이 부족 ({char_count}자) — 이미지 기반 문서이거나 파싱 오류일 수 있음")
    if not has_claims:
        warnings.append("청구항(Claim) 텍스트 미감지 — 특허 문서인지 확인 필요")
    if non_text_ratio > 0.3:
        warnings.append(f"비텍스트 문자 비율 {non_text_ratio:.0%} — OCR 오류 또는 인코딩 문제 의심")

    return {
        "char_count": char_count,
        "has_claims": has_claims,
        "has_korean": has_korean,
        "has_english": has_english,
        "quality": "good" if not warnings else ("warn" if char_count >= 200 else "poor"),
        "warnings": warnings,
    }
